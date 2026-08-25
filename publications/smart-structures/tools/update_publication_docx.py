#!/usr/bin/env python3
"""Prepare the Smart Structures DOCX as a reproducible release candidate.

The script performs surgical content/metadata fixes, marks table rows as
non-splittable, adds live bibliography hyperlinks, and embeds the three OFL
font families used by the document according to ECMA-376 Part 1, section
17.8.1.
"""

from __future__ import annotations

import argparse
import copy
import datetime as dt
import re
import uuid
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.shared import Pt, RGBColor
from lxml import etree


W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
CT_NS = "http://schemas.openxmlformats.org/package/2006/content-types"
DC_NS = "http://purl.org/dc/elements/1.1/"


def _set_run_font(run, family: str, size_pt: float | None = None) -> None:
    run.font.name = family
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    rpr = run._element.get_or_add_rPr()
    fonts = rpr.find(qn("w:rFonts"))
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.insert(0, fonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{attr}"), family)


def _replace_paragraph_text(paragraph, text: str) -> None:
    template_rpr = None
    if paragraph.runs:
        existing = paragraph.runs[0]._element.find(qn("w:rPr"))
        if existing is not None:
            template_rpr = copy.deepcopy(existing)
    paragraph.clear()
    run = paragraph.add_run(text)
    if template_rpr is not None:
        current = run._element.find(qn("w:rPr"))
        if current is not None:
            run._element.remove(current)
        run._element.insert(0, template_rpr)


def _add_external_hyperlink(
    paragraph,
    text: str,
    url: str,
    *,
    rtl_text: bool = False,
    language: str = "en-US",
):
    rel_id = paragraph.part.relate_to(
        url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)

    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "136B8A")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rtl = OxmlElement("w:rtl")
    rtl.set(qn("w:val"), "1" if rtl_text else "0")
    lang = OxmlElement("w:lang")
    lang.set(qn("w:val"), language)
    rpr.extend((color, underline, rtl, lang))
    run.append(rpr)
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    return hyperlink


def _replace_reference(paragraph, citation: str, doi_url: str | None = None) -> None:
    paragraph.clear()
    if doi_url and citation.endswith(doi_url):
        prefix = citation[: -len(doi_url)]
        paragraph.add_run(prefix)
        _add_external_hyperlink(
            paragraph,
            "doi:" + doi_url.removeprefix("https://doi.org/"),
            doi_url,
        )
    else:
        paragraph.add_run(citation)


def _set_title_author(paragraph) -> None:
    paragraph.clear()
    fa = paragraph.add_run("یوسف بهرام بیگی")
    _set_run_font(fa, "Sahel", 16)
    fa.bold = True
    fa.font.color.rgb = RGBColor(0x07, 0x1A, 0x2D)
    fa._element.get_or_add_rPr().append(OxmlElement("w:rtl"))
    fa.add_break()
    en = paragraph.add_run("Yousef Bahrambeigi")
    _set_run_font(en, "Shabnam", 10)
    en.bold = True
    en.font.color.rgb = RGBColor(0x62, 0x78, 0x8A)
    rtl = OxmlElement("w:rtl")
    rtl.set(qn("w:val"), "0")
    en._element.get_or_add_rPr().append(rtl)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _add_release_notice(doc: Document, final_identifier: str | None) -> None:
    cell = doc.tables[0].cell(0, 0)
    license_p = cell.add_paragraph()
    license_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    ppr = license_p._p.get_or_add_pPr()
    ppr.append(OxmlElement("w:bidi"))
    run = license_p.add_run(
        "© ۲۰۲۶ یوسف بهرام بیگی. متن، شکل‌های اصیل، جدول‌ها و نمونه‌کدها، "
        "جز در مواردی که خلاف آن تصریح شده است، تحت مجوز CC BY 4.0 منتشر می‌شوند. "
        "مطالب اشخاص ثالث تابع حقوق صاحبان اصلی‌اند. "
    )
    _set_run_font(run, "Vazirmatn", 8.5)
    run.font.color.rgb = RGBColor(0x24, 0x37, 0x44)
    run._element.get_or_add_rPr().append(OxmlElement("w:rtl"))
    link_p = cell.add_paragraph()
    link_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _add_external_hyperlink(
        link_p,
        "Full CC BY 4.0 license",
        "https://creativecommons.org/licenses/by/4.0/",
    )

    status_p = cell.add_paragraph()
    status_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    status_p._p.get_or_add_pPr().append(OxmlElement("w:bidi"))
    if final_identifier:
        status = f"نسخه ۱.۰.۰ | شناسه رسمی: {final_identifier}"
    else:
        status = (
            "نامزد انتشار ۱.۰.۰ (RC1) | شناسه رسمی پیش از انتشار نهایی درج می‌شود."
        )
    run = status_p.add_run(status)
    _set_run_font(run, "Shabnam", 8.5)
    run.bold = True
    run.font.color.rgb = RGBColor(0xD5, 0xA7, 0x4E)
    run._element.get_or_add_rPr().append(OxmlElement("w:rtl"))


def _set_table_row_rules(doc: Document) -> None:
    for table in doc.tables:
        for row in table.rows:
            trpr = row._tr.get_or_add_trPr()
            if trpr.find(qn("w:cantSplit")) is None:
                trpr.append(OxmlElement("w:cantSplit"))
        if len(table.rows) > 1:
            trpr = table.rows[0]._tr.get_or_add_trPr()
            if trpr.find(qn("w:tblHeader")) is None:
                trpr.append(OxmlElement("w:tblHeader"))


def _patch_content(doc: Document, final_identifier: str | None) -> None:
    if len(doc.paragraphs) < 349 or len(doc.tables) < 33:
        raise ValueError("Unexpected document structure; refusing an unsafe patch")

    _set_title_author(doc.paragraphs[7])
    _replace_paragraph_text(
        doc.paragraphs[8],
        "ویرایش دانشگاهی نخست | "
        + ("نسخه ۱.۰.۰" if final_identifier else "نامزد انتشار ۱.۰.۰ (RC1)")
        + " | ۱۴۰۵",
    )
    _add_release_notice(doc, final_identifier)

    _replace_paragraph_text(
        doc.paragraphs[34],
        doc.paragraphs[34].text.rstrip(" .")
        + " این چارچوب با مرورهای کلاسیک کنترل سازه نیز هم‌خوان است [۳، ۴].",
    )
    _replace_paragraph_text(
        doc.paragraphs[137],
        doc.paragraphs[137].text.rstrip(" .")
        + " مرورهای تخصصی سامانه‌های اتلاف انرژی و میراگرهای تکمیلی نیز بر همین انتخاب مبتنی بر عملکرد تأکید دارند [۵، ۱۱].",
    )
    _replace_paragraph_text(
        doc.paragraphs[287],
        doc.paragraphs[287].text.replace("[۶، ۸]", "[۶، ۷، ۸]"),
    )

    references = {
        334: (
            "Cheng, F. Y., Jiang, H., & Lou, K. (2008). Smart Structures: Innovative Systems for Seismic Response Control. CRC Press. ISBN 978-0-8493-8532-2.",
            None,
        ),
        335: (
            "ASCE/SEI. (2022). Minimum Design Loads and Associated Criteria for Buildings and Other Structures (ASCE/SEI 7-22). American Society of Civil Engineers. https://doi.org/10.1061/9780784415788",
            "https://doi.org/10.1061/9780784415788",
        ),
        336: (
            "Spencer, B. F., Jr., & Nagarajaiah, S. (2003). State of the Art of Structural Control. Journal of Structural Engineering, 129(7), 845–856. https://doi.org/10.1061/(ASCE)0733-9445(2003)129:7(845)",
            "https://doi.org/10.1061/(ASCE)0733-9445(2003)129:7(845)",
        ),
        337: (
            "Housner, G. W., Bergman, L. A., Caughey, T. K., Chassiakos, A. G., Claus, R. O., Masri, S. F., Skelton, R. E., Soong, T. T., Spencer, B. F., & Yao, J. T. P. (1997). Structural Control: Past, Present, and Future. Journal of Engineering Mechanics, 123(9), 897–971. https://doi.org/10.1061/(ASCE)0733-9399(1997)123:9(897)",
            "https://doi.org/10.1061/(ASCE)0733-9399(1997)123:9(897)",
        ),
        338: (
            "Soong, T. T., & Dargush, G. F. (1997). Passive Energy Dissipation Systems in Structural Engineering. Wiley. ISBN 978-0-471-96821-4.",
            None,
        ),
        339: (
            "Nasim, M., Rajabifard, A., Chen, Y., & Samali, B. (2026). A demonstration of a digital twin framework for structural health monitoring: Application to bridge infrastructures. Journal of Infrastructure Intelligence and Resilience, 5(1), 100184. https://doi.org/10.1016/j.iintel.2025.100184",
            "https://doi.org/10.1016/j.iintel.2025.100184",
        ),
        340: (
            "Sun, Z., Jayasinghe, S., Sidiq, A., Shahrivar, F., Mahmoodian, M., & Setunge, S. (2025). Approach Towards the Development of Digital Twin for Structural Health Monitoring of Civil Infrastructure: A Comprehensive Review. Sensors, 25(1), 59. https://doi.org/10.3390/s25010059",
            "https://doi.org/10.3390/s25010059",
        ),
        341: (
            "Zhai, G., Yao, Z., Wang, D., Zhang, A. A., Spencer, B. F., & Xu, Y. (2026). Autonomous post-earthquake structural assessment based on bidirectional graphics-based digital twin (Bi-GBDT) with physical and visual realism. Advanced Engineering Informatics, 69, 103971. https://doi.org/10.1016/j.aei.2025.103971",
            "https://doi.org/10.1016/j.aei.2025.103971",
        ),
        342: (
            "Hu, S., Guo, T., Alam, M. S., Koetaka, Y., Ghafoori, E., & Karavasilis, T. L. (2025). Machine learning in earthquake engineering: A review on recent progress and future trends in seismic performance evaluation and design. Engineering Structures, 340, 120721. https://doi.org/10.1016/j.engstruct.2025.120721",
            "https://doi.org/10.1016/j.engstruct.2025.120721",
        ),
        343: (
            "Ghanemi, N. E., Abdeddaim, M., Ounis, A., & Basili, M. (2025). Neural network based active control of base isolated structure considering isolator nonlinearity. Frontiers in Built Environment, 11, 1630131. https://doi.org/10.3389/fbuil.2025.1630131",
            "https://doi.org/10.3389/fbuil.2025.1630131",
        ),
        344: (
            "Katsimpini, P., Papagiannopoulos, G., & Hatzigeorgiou, G. (2025). A Thorough Examination of Innovative Supplementary Dampers Aimed at Enhancing the Seismic Behavior of Structural Systems. Applied Sciences, 15(3), 1226. https://doi.org/10.3390/app15031226",
            "https://doi.org/10.3390/app15031226",
        ),
    }
    for index, (citation, doi_url) in references.items():
        _replace_reference(doc.paragraphs[index], citation, doi_url)

    citation_status = (
        f"شناسه رسمی: {final_identifier}." if final_identifier else "شناسه رسمی در انتشار نهایی افزوده می‌شود."
    )
    _replace_paragraph_text(
        doc.paragraphs[346],
        "بهرام بیگی، یوسف. (۱۴۰۵/۲۰۲۶). سازه‌های هوشمند و کنترل پاسخ لرزه‌ای: "
        "از جداسازی پایه و میرایی تا کنترل فعال، دوقلوی دیجیتال و هوش سازه‌ای. "
        + ("ویرایش دانشگاهی نخست، نسخه ۱.۰.۰. " if final_identifier else "ویرایش دانشگاهی نخست، نامزد انتشار ۱.۰.۰ (RC1). ")
        + citation_status
        + " مجوز CC BY 4.0.",
    )

    # Keep the closing statement with the citation page.  The source document
    # forced it onto an otherwise empty page, which became wasteful once the
    # bibliography was corrected and expanded.
    closing_ppr = doc.paragraphs[347]._p.get_or_add_pPr()
    page_break_before = closing_ppr.find(qn("w:pageBreakBefore"))
    if page_break_before is not None:
        closing_ppr.remove(page_break_before)
    spacing = closing_ppr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        closing_ppr.append(spacing)
    spacing.set(qn("w:before"), "800")

    _set_table_row_rules(doc)


def _set_core_properties(doc: Document, final_identifier: str | None) -> None:
    props = doc.core_properties
    props.title = "سازه‌های هوشمند و کنترل پاسخ لرزه‌ای"
    props.subject = "Smart Structures and Seismic Response Control"
    props.author = "Yousef Bahrambeigi (یوسف بهرام بیگی)"
    props.last_modified_by = "Yousef Bahrambeigi"
    props.keywords = (
        "smart structures; seismic control; base isolation; damping; active control; "
        "digital twin; structural health monitoring; CC BY 4.0"
    )
    status = "version 1.0.0" if final_identifier else "release candidate 1 for version 1.0.0"
    identifier = f" Identifier: {final_identifier}." if final_identifier else ""
    props.comments = (
        "Persian academic book; " + status + ". Licensed under CC BY 4.0." + identifier
    )
    props.category = "Academic book; Civil Engineering; Smart Structures"
    props.created = dt.datetime(2026, 8, 25, 0, 0, 0, tzinfo=dt.timezone.utc)
    props.modified = dt.datetime.now(dt.timezone.utc).replace(microsecond=0)
    props.revision = 2


def _font_key_and_bytes(font_bytes: bytes) -> tuple[str, bytes]:
    if len(font_bytes) < 32:
        raise ValueError("Font file is unexpectedly short")
    guid = uuid.uuid4()
    key = guid.bytes[::-1]
    obfuscated = bytearray(font_bytes)
    for offset in range(32):
        obfuscated[offset] ^= key[offset % 16]
    return "{" + str(guid).upper() + "}", bytes(obfuscated)


def _next_rel_id(root) -> str:
    used = set()
    for rel in root.findall(f"{{{REL_NS}}}Relationship"):
        match = re.fullmatch(r"rId(\d+)", rel.get("Id", ""))
        if match:
            used.add(int(match.group(1)))
    value = 1
    while value in used:
        value += 1
    return f"rId{value}"


def _embed_fonts(docx_path: Path, fonts: dict[str, dict[str, Path]]) -> None:
    with zipfile.ZipFile(docx_path, "r") as archive:
        parts = {name: archive.read(name) for name in archive.namelist()}

    parser = etree.XMLParser(remove_blank_text=False)
    font_table = etree.fromstring(parts["word/fontTable.xml"], parser)
    rels_name = "word/_rels/fontTable.xml.rels"
    if rels_name in parts:
        rels = etree.fromstring(parts[rels_name], parser)
    else:
        rels = etree.Element(f"{{{REL_NS}}}Relationships", nsmap={None: REL_NS})

    for existing in list(rels.findall(f"{{{REL_NS}}}Relationship")):
        if existing.get("Type", "").endswith("/font"):
            rels.remove(existing)

    for family, variants in fonts.items():
        font_node = font_table.find(f"{{{W_NS}}}font[@{{{W_NS}}}name='{family}']")
        if font_node is None:
            font_node = etree.SubElement(font_table, f"{{{W_NS}}}font")
            font_node.set(f"{{{W_NS}}}name", family)
        for child_name in ("embedRegular", "embedBold", "embedItalic", "embedBoldItalic"):
            for child in list(font_node.findall(f"{{{W_NS}}}{child_name}")):
                font_node.remove(child)

        for variant, font_path in variants.items():
            rel_id = _next_rel_id(rels)
            file_index = 1
            while f"word/fonts/font{file_index}.odttf" in parts:
                file_index += 1
            target = f"fonts/font{file_index}.odttf"
            rel = etree.SubElement(rels, f"{{{REL_NS}}}Relationship")
            rel.set("Id", rel_id)
            rel.set(
                "Type",
                "http://schemas.openxmlformats.org/officeDocument/2006/relationships/font",
            )
            rel.set("Target", target)
            font_key, embedded = _font_key_and_bytes(font_path.read_bytes())
            parts[f"word/{target}"] = embedded
            element_name = "embedRegular" if variant == "regular" else "embedBold"
            embed = etree.SubElement(font_node, f"{{{W_NS}}}{element_name}")
            embed.set(f"{{{R_NS}}}id", rel_id)
            embed.set(f"{{{W_NS}}}fontKey", font_key)

    parts["word/fontTable.xml"] = etree.tostring(
        font_table, xml_declaration=True, encoding="UTF-8", standalone=True
    )
    parts[rels_name] = etree.tostring(
        rels, xml_declaration=True, encoding="UTF-8", standalone=True
    )

    settings = etree.fromstring(parts["word/settings.xml"], parser)
    for tag in ("embedTrueTypeFonts", "saveSubsetFonts"):
        if settings.find(f"{{{W_NS}}}{tag}") is None:
            settings.insert(0, etree.Element(f"{{{W_NS}}}{tag}"))
    parts["word/settings.xml"] = etree.tostring(
        settings, xml_declaration=True, encoding="UTF-8", standalone=True
    )

    content_types = etree.fromstring(parts["[Content_Types].xml"], parser)
    has_odttf = any(
        node.get("Extension") == "odttf"
        for node in content_types.findall(f"{{{CT_NS}}}Default")
    )
    if not has_odttf:
        default = etree.SubElement(content_types, f"{{{CT_NS}}}Default")
        default.set("Extension", "odttf")
        default.set(
            "ContentType",
            "application/vnd.openxmlformats-officedocument.obfuscatedFont",
        )
    parts["[Content_Types].xml"] = etree.tostring(
        content_types, xml_declaration=True, encoding="UTF-8", standalone=True
    )

    core = etree.fromstring(parts["docProps/core.xml"], parser)
    language = core.find(f"{{{DC_NS}}}language")
    if language is None:
        language = etree.SubElement(core, f"{{{DC_NS}}}language")
    language.text = "fa-IR"
    parts["docProps/core.xml"] = etree.tostring(
        core, xml_declaration=True, encoding="UTF-8", standalone=True
    )

    text_content = []
    for paragraph in Document(docx_path).paragraphs:
        text_content.append(paragraph.text)
    joined = "\n".join(text_content)
    words = len(re.findall(r"\S+", joined))
    characters = len(joined.replace("\n", ""))
    app = etree.fromstring(parts["docProps/app.xml"], parser)
    app_ns = app.nsmap.get(None)
    values = {
        "Pages": "36",
        "Words": str(words),
        "Characters": str(characters),
        "CharactersWithSpaces": str(len(joined)),
        "Paragraphs": str(len(text_content)),
        "Application": "Microsoft Word / OOXML publication pipeline",
    }
    for tag, value in values.items():
        node = app.find(f"{{{app_ns}}}{tag}")
        if node is not None:
            node.text = value
    parts["docProps/app.xml"] = etree.tostring(
        app, xml_declaration=True, encoding="UTF-8", standalone=True
    )

    temp_path = docx_path.with_suffix(".embedded.tmp")
    with zipfile.ZipFile(temp_path, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        for name in sorted(parts):
            archive.writestr(name, parts[name])
    temp_path.replace(docx_path)


def _verify_embedded_fonts(docx_path: Path, expected_count: int) -> None:
    with zipfile.ZipFile(docx_path, "r") as archive:
        names = [name for name in archive.namelist() if name.endswith(".odttf")]
        if len(names) != expected_count:
            raise ValueError(f"Expected {expected_count} embedded fonts, found {len(names)}")
        font_table = etree.fromstring(archive.read("word/fontTable.xml"))
        rels = etree.fromstring(archive.read("word/_rels/fontTable.xml.rels"))
        rel_targets = {
            rel.get("Id"): rel.get("Target")
            for rel in rels.findall(f"{{{REL_NS}}}Relationship")
            if rel.get("Type", "").endswith("/font")
        }
        verified = 0
        for family in ("Shabnam", "Sahel", "Vazirmatn"):
            node = font_table.find(f"{{{W_NS}}}font[@{{{W_NS}}}name='{family}']")
            if node is None:
                raise ValueError(f"Missing font table entry for {family}")
            for embed_name in ("embedRegular", "embedBold"):
                embed = node.find(f"{{{W_NS}}}{embed_name}")
                if embed is None:
                    raise ValueError(f"Missing {embed_name} for {family}")
                rel_id = embed.get(f"{{{R_NS}}}id")
                key_text = embed.get(f"{{{W_NS}}}fontKey").strip("{}")
                key = bytes.fromhex(key_text.replace("-", ""))[::-1]
                embedded = bytearray(archive.read("word/" + rel_targets[rel_id]))
                for offset in range(32):
                    embedded[offset] ^= key[offset % 16]
                if bytes(embedded[:4]) not in (b"\x00\x01\x00\x00", b"OTTO", b"true", b"ttcf"):
                    raise ValueError(f"Embedded {family} {embed_name} did not decode as a font")
                verified += 1
        if verified != expected_count:
            raise ValueError(f"Verified {verified} embedded fonts, expected {expected_count}")


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser()
    parser.add_argument("--input", type=Path, required=True)
    parser.add_argument("--output", type=Path, required=True)
    parser.add_argument("--shabnam-dir", type=Path, required=True)
    parser.add_argument("--sahel-dir", type=Path, required=True)
    parser.add_argument("--vazirmatn-dir", type=Path, required=True)
    parser.add_argument("--identifier", help="Final DOI or ISBN; omit for RC1")
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    doc = Document(args.input)
    _patch_content(doc, args.identifier)
    _set_core_properties(doc, args.identifier)
    args.output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(args.output)

    fonts = {
        "Shabnam": {
            "regular": args.shabnam_dir / "Shabnam.ttf",
            "bold": args.shabnam_dir / "Shabnam-Bold.ttf",
        },
        "Sahel": {
            "regular": args.sahel_dir / "Sahel.ttf",
            "bold": args.sahel_dir / "Sahel-Bold.ttf",
        },
        "Vazirmatn": {
            "regular": args.vazirmatn_dir / "Vazirmatn-Regular.ttf",
            "bold": args.vazirmatn_dir / "Vazirmatn-Bold.ttf",
        },
    }
    _embed_fonts(args.output, fonts)
    _verify_embedded_fonts(args.output, expected_count=6)
    with zipfile.ZipFile(args.output, "r") as archive:
        bad = archive.testzip()
        if bad:
            raise ValueError(f"Corrupt ZIP member: {bad}")
    print(f"Prepared {args.output}")


if __name__ == "__main__":
    main()
